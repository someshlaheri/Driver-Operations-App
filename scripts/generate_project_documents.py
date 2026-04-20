from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


ROOT = Path(__file__).resolve().parents[1]
DOCS_DIR = ROOT / "docs"
DOCX_PATH = DOCS_DIR / "Driver-Operations-App-Documentation-Formatted.docx"
PDF_PATH = DOCS_DIR / "Driver-Operations-App-Documentation-Formatted.pdf"


ACCENT = RGBColor(0x1F, 0x5A, 0xA6)
TEXT_DARK = RGBColor(0x22, 0x2B, 0x38)
TEXT_SOFT = RGBColor(0x5C, 0x66, 0x73)

PDF_ACCENT = colors.HexColor("#1F5AA6")
PDF_DARK = colors.HexColor("#222B38")
PDF_SOFT = colors.HexColor("#5C6673")
PDF_BORDER = colors.HexColor("#CAD3DD")
PDF_FILL = colors.HexColor("#EEF4FB")


@dataclass(frozen=True)
class Section:
    title: str
    paragraphs: Sequence[str] = ()
    bullets: Sequence[str] = ()
    table: Sequence[Sequence[str]] | None = None
    page_break_before: bool = False


def resolve_writable_path(path: Path) -> Path:
    candidate = path
    counter = 2

    while True:
        try:
            if candidate.exists():
                with open(candidate, "ab"):
                    pass
            return candidate
        except PermissionError:
            candidate = path.with_name(f"{path.stem}-{counter}{path.suffix}")
            counter += 1


def set_cell_shading(cell, fill: str) -> None:
    cell_properties = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    cell_properties.append(shading)


def set_document_margins(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)


def ensure_styles(document: Document) -> None:
    styles = document.styles

    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = TEXT_DARK

    for style_name, size, color in [
        ("Title", 22, ACCENT),
        ("Heading 1", 16, ACCENT),
        ("Heading 2", 12.5, TEXT_DARK),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color

    if "Cover Subtitle" not in styles:
        subtitle = styles.add_style("Cover Subtitle", WD_STYLE_TYPE.PARAGRAPH)
        subtitle.font.name = "Calibri"
        subtitle.font.size = Pt(11)
        subtitle.font.color.rgb = TEXT_SOFT


def add_cover_page(document: Document) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Driver Operations App")
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = ACCENT

    p = document.add_paragraph(style="Cover Subtitle")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Functional and Technical Documentation")

    p = document.add_paragraph(style="Cover Subtitle")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Full-stack demo system for driver scheduling, shift tracking, vehicle inspections, and admin oversight")

    document.add_paragraph("")

    snapshot = [
        ["Project Type", "Internal operations web application"],
        ["Frontend", "React, TypeScript, Vite"],
        ["Backend", "Node.js, Express, TypeScript"],
        ["Database", "PostgreSQL with Prisma"],
        ["Authentication", "JWT with bcrypt password hashing"],
        ["Primary Roles", "Driver, Admin"],
    ]
    add_docx_table(document, ["Item", "Details"], snapshot, "D9E8F7")

    document.add_paragraph("")
    p = document.add_paragraph(style="Cover Subtitle")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Generated from the current repository implementation")

    document.add_page_break()


def add_docx_table(document: Document, headers: Sequence[str], rows: Sequence[Sequence[str]], header_fill: str) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        header_cells[index].text = header
        set_cell_shading(header_cells[index], header_fill)
        for paragraph in header_cells[index].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = TEXT_DARK

    for row in rows:
        row_cells = table.add_row().cells
        for index, value in enumerate(row):
            row_cells[index].text = value

    document.add_paragraph("")


def add_docx_section(document: Document, section: Section) -> None:
    if section.page_break_before:
      document.add_page_break()

    document.add_heading(section.title, level=1)

    for paragraph in section.paragraphs:
        document.add_paragraph(paragraph)

    for bullet in section.bullets:
        document.add_paragraph(bullet, style="List Bullet")

    if section.table:
        headers = section.table[0]
        rows = section.table[1:]
        add_docx_table(document, headers, rows, "D9E8F7")


def build_docx(sections: Iterable[Section]) -> None:
    document = Document()
    set_document_margins(document)
    ensure_styles(document)
    add_cover_page(document)

    for section in sections:
        add_docx_section(document, section)

    output_path = resolve_writable_path(DOCX_PATH)
    document.save(output_path)
    return output_path


def get_pdf_styles():
    styles = getSampleStyleSheet()

    styles["Title"].fontName = "Helvetica-Bold"
    styles["Title"].fontSize = 22
    styles["Title"].textColor = PDF_ACCENT
    styles["Title"].alignment = TA_CENTER
    styles["Title"].spaceAfter = 12

    styles["Heading1"].fontName = "Helvetica-Bold"
    styles["Heading1"].fontSize = 15
    styles["Heading1"].textColor = PDF_ACCENT
    styles["Heading1"].spaceBefore = 12
    styles["Heading1"].spaceAfter = 8

    styles["BodyText"].fontName = "Helvetica"
    styles["BodyText"].fontSize = 10
    styles["BodyText"].leading = 14
    styles["BodyText"].textColor = PDF_DARK
    styles["BodyText"].spaceAfter = 6

    styles.add(
        ParagraphStyle(
            name="CoverSub",
            parent=styles["BodyText"],
            alignment=TA_CENTER,
            textColor=PDF_SOFT,
            spaceAfter=6,
        )
    )
    styles.add(
        ParagraphStyle(
            name="BulletBody",
            parent=styles["BodyText"],
            leftIndent=14,
            bulletIndent=4,
            spaceAfter=4,
        )
    )
    return styles


def pdf_table(data: Sequence[Sequence[str]], col_widths: Sequence[float] | None = None) -> Table:
    table = Table(data, colWidths=col_widths, repeatRows=1)
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), PDF_FILL),
                ("TEXTCOLOR", (0, 0), (-1, 0), PDF_DARK),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTNAME", (0, 1), (-1, -1), "Helvetica"),
                ("FONTSIZE", (0, 0), (-1, -1), 9.5),
                ("GRID", (0, 0), (-1, -1), 0.5, PDF_BORDER),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ]
        )
    )
    return table


def build_pdf(sections: Iterable[Section]) -> None:
    styles = get_pdf_styles()
    story = []

    story.append(Paragraph("Driver Operations App", styles["Title"]))
    story.append(Paragraph("Functional and Technical Documentation", styles["CoverSub"]))
    story.append(
        Paragraph(
            "Full-stack demo system for driver scheduling, shift tracking, vehicle inspections, and admin oversight",
            styles["CoverSub"],
        )
    )
    story.append(Spacer(1, 0.2 * inch))

    story.append(
        pdf_table(
            [
                ["Item", "Details"],
                ["Project Type", "Internal operations web application"],
                ["Frontend", "React, TypeScript, Vite"],
                ["Backend", "Node.js, Express, TypeScript"],
                ["Database", "PostgreSQL with Prisma"],
                ["Authentication", "JWT with bcrypt password hashing"],
                ["Primary Roles", "Driver, Admin"],
            ],
            col_widths=[1.5 * inch, 4.8 * inch],
        )
    )
    story.append(Spacer(1, 0.25 * inch))
    story.append(Paragraph("Generated from the current repository implementation", styles["CoverSub"]))
    story.append(PageBreak())

    for section in sections:
        if section.page_break_before:
            story.append(PageBreak())

        story.append(Paragraph(section.title, styles["Heading1"]))

        for paragraph in section.paragraphs:
            story.append(Paragraph(paragraph, styles["BodyText"]))

        for bullet in section.bullets:
            story.append(Paragraph(bullet, styles["BulletBody"], bulletText="\u2022"))

        if section.table:
            story.append(pdf_table(section.table))
            story.append(Spacer(1, 0.12 * inch))

    output_path = resolve_writable_path(PDF_PATH)

    document = SimpleDocTemplate(
        str(output_path),
        pagesize=A4,
        rightMargin=0.6 * inch,
        leftMargin=0.6 * inch,
        topMargin=0.55 * inch,
        bottomMargin=0.55 * inch,
        title="Driver Operations App - Functional and Technical Documentation",
    )
    document.build(story)
    return output_path


def get_sections() -> list[Section]:
    return [
        Section(
            "1. Executive Summary",
            paragraphs=[
                "The Driver Operations App is a production-style internal workflow platform built to manage driver scheduling, live shift execution, daily vehicle inspections, and admin-side review.",
                "The system models a realistic logistics workflow rather than a simple CRUD exercise. Drivers can self-register, authenticate, claim shifts, start work, submit inspection evidence with photo upload, and complete shifts. Admin users can register with an internal code, oversee coverage, create shifts, and review inspection outcomes from a dedicated control dashboard.",
            ],
        ),
        Section(
            "2. Business Goals",
            bullets=[
                "Digitize driver scheduling and inspection workflows in a single web system.",
                "Reduce manual coordination by allowing drivers to claim shifts directly.",
                "Enforce inspection completion before a shift can be closed.",
                "Provide admins with live operational visibility across drivers, shifts, and inspection compliance.",
                "Maintain persistent records for claims, inspections, review status, and shift history.",
            ],
        ),
        Section(
            "3. User Roles and Responsibilities",
            table=[
                ["Role", "Responsibilities", "Key Access"],
                ["Driver", "Can self-register, claim shifts, start work, submit inspections, and end shifts", "Driver dashboard and driver APIs"],
                ["Admin", "Registers with internal code, monitors operations, creates shifts, reviews inspections, and tracks activity", "Admin dashboard and admin APIs"],
            ],
        ),
        Section(
            "4. Functional Modules",
            table=[
                ["Module", "Description", "Primary Outcome"],
                ["Authentication", "Driver/admin registration plus login with JWT session restoration", "Secure role-based access"],
                ["Shift Discovery", "Upcoming 7-day shift listing with capacity tracking", "Drivers can self-assign work"],
                ["Shift Claiming", "Drivers claim available shifts with duplicate and capacity checks", "Prevents overbooking"],
                ["Shift Lifecycle", "Claimed -> Started -> Completed state transitions", "Operational tracking"],
                ["Inspection Submission", "Vehicle inspection form with mandatory photo upload", "Compliance evidence"],
                ["Admin Dashboard", "Operational summary, driver status, shifts, inspections", "Centralized oversight"],
                ["Inspection Review", "Approve, flag, or reset inspection review state with notes", "Review and quality control"],
            ],
        ),
        Section(
            "5. End-to-End Workflow",
            paragraphs=[
                "The driver workflow is intentionally sequential and enforced on the backend. A driver first claims a shift, then starts it, then submits an inspection, and only after that can end the shift.",
                "The admin workflow centers on visibility and control. Admin users review upcoming shift coverage, monitor active drivers, create shifts for open demand, and update inspection review outcomes.",
            ],
            bullets=[
                "Driver workflow: Login -> Claim Shift -> Start Shift -> Submit Inspection -> End Shift",
                "Admin workflow: Login -> Review Dashboard -> Create Shift -> Review Inspection -> Update Status",
            ],
        ),
        Section(
            "6. Business Rules",
            bullets=[
                "A driver cannot claim the same shift more than once.",
                "A shift cannot be claimed once capacity is full.",
                "Only a claimed shift can be started.",
                "Only a started shift can receive an inspection.",
                "Only one inspection is allowed per shift claim.",
                "A started shift cannot be completed until an inspection exists.",
                "Inspection photo upload is mandatory.",
                "Admin-only endpoints are protected from driver access, and vice versa.",
            ],
        ),
        Section(
            "7. Technology Stack",
            table=[
                ["Layer", "Technology", "Purpose"],
                ["Frontend", "React 19, TypeScript, Vite, React Router", "User interface and routing"],
                ["Backend", "Node.js, Express 5, TypeScript", "REST API and business logic"],
                ["Database", "PostgreSQL 18", "Persistent operational data"],
                ["ORM", "Prisma 7 with Prisma Postgres adapter", "Schema modeling and queries"],
                ["Authentication", "JWT, bcryptjs", "Session and credential security"],
                ["Validation", "Zod", "Input validation and error shaping"],
                ["Uploads", "Multer", "Multipart image handling"],
            ],
        ),
        Section(
            "8. System Architecture",
            paragraphs=[
                "The project uses a monorepo structure with separate client and server workspaces. The frontend communicates with the backend via REST endpoints under /api, and the backend persists all core entities to PostgreSQL.",
                "In production, the Express server serves the compiled React application from the client build output, allowing a single backend service to host both the frontend and the API.",
            ],
            table=[
                ["Component", "Responsibility"],
                ["client/", "React frontend, route protection, dashboard UI, forms, API calls"],
                ["server/src/routes", "Feature-based API endpoints for auth, driver, and admin flows"],
                ["server/src/lib", "Reusable auth and database client utilities"],
                ["server/prisma", "Schema definition and seed logic"],
                ["docs/", "Human-readable project documentation"],
            ],
        ),
        Section(
            "9. Database Model",
            table=[
                ["Entity", "Purpose", "Key Constraints"],
                ["User", "Stores admins and drivers", "Unique email, role enum"],
                ["Shift", "Stores scheduled work slots", "Indexed by date"],
                ["ShiftClaim", "Connects drivers to shifts", "Unique shiftId + driverId"],
                ["VehicleInspection", "Stores inspection details and photo evidence", "Unique claimId"],
            ],
            paragraphs=[
                "The data model uses relational integrity to enforce core workflow constraints. VehicleInspection has a one-to-one relationship with ShiftClaim, ensuring only one inspection per claimed shift. ShiftClaim connects a driver to a shift and stores lifecycle timestamps such as claimedAt, startedAt, and endedAt.",
                "Inspection images are stored directly in the database as binary data for simpler demo deployment. This avoids local file storage complexity and keeps the system self-contained.",
            ],
        ),
        Section(
            "10. Authentication and Authorization Design",
            table=[
                ["Capability", "Implementation"],
                ["Registration", "Driver self-registration and admin registration with internal code"],
                ["Password storage", "bcryptjs hashing"],
                ["Session token", "JWT signed with server secret"],
                ["Session duration", "7 days"],
                ["Client token storage", "localStorage"],
                ["Auth middleware", "Bearer token validation via requireAuth"],
                ["Role enforcement", "requireRole('DRIVER') and requireRole('ADMIN')"],
            ],
        ),
        Section(
            "11. API Inventory",
            page_break_before=True,
            table=[
                ["Method", "Path", "Purpose"],
                ["GET", "/api/health", "Service health check"],
                ["POST", "/api/auth/register", "Register a new driver or admin account"],
                ["POST", "/api/auth/login", "Authenticate user and return JWT"],
                ["GET", "/api/auth/me", "Restore current authenticated user"],
                ["GET", "/api/driver/dashboard", "Return driver shifts, claims, and metrics"],
                ["POST", "/api/driver/shifts/:shiftId/claim", "Claim a shift"],
                ["POST", "/api/driver/claims/:claimId/start", "Start a claimed shift"],
                ["POST", "/api/driver/claims/:claimId/inspection", "Submit vehicle inspection with photo"],
                ["POST", "/api/driver/claims/:claimId/end", "Complete a started shift"],
                ["GET", "/api/admin/dashboard", "Return admin dashboard data"],
                ["POST", "/api/admin/shifts", "Create a new shift"],
                ["PATCH", "/api/admin/shifts/:shiftId", "Update an existing shift"],
                ["PATCH", "/api/admin/inspections/:inspectionId/review", "Update inspection review status and notes"],
                ["GET", "/api/inspections/:inspectionId/photo", "Stream inspection image from database"],
            ],
        ),
        Section(
            "12. Validation and Data Quality Controls",
            bullets=[
                "Login requires a valid email format and non-empty password.",
                "Shift capacity is constrained between 1 and 10.",
                "Shift times must follow HH:MM format.",
                "Mileage must be a non-negative integer.",
                "Fuel percentage must stay between 0 and 100.",
                "Cleanliness score must stay between 1 and 5.",
                "Inspection photo uploads are limited to 5 MB.",
                "Review notes and issue descriptions are length-constrained.",
            ],
        ),
        Section(
            "13. Frontend Design Notes",
            paragraphs=[
                "The frontend is organized around role-based entry points. After login, users are redirected to either the driver workspace or the admin workspace based on the role stored in the token and returned from /api/auth/me.",
                "The driver screen emphasizes operational action flow, while the admin screen emphasizes oversight, coverage monitoring, and review actions. The browser title updates per page and the app uses a custom SVG favicon for a more polished demo experience.",
            ],
            table=[
                ["Page", "Purpose"],
                ["LoginPage", "Role-aware sign-in entry point"],
                ["RegisterPage", "Driver/admin registration with admin-code protection"],
                ["DriverDashboard", "Shift claiming, live shift execution, inspection submission"],
                ["AdminDashboard", "Shift creation, operational metrics, inspection review"],
            ],
        ),
        Section(
            "14. Deployment and Runtime Setup",
            table=[
                ["Setting", "Value / Notes"],
                ["Frontend dev port", "5173"],
                ["Backend dev port", "4000"],
                ["Database port", "5432"],
                ["Production startup", "npm run build && npm start"],
                ["Backend entrypoint", "server/dist/src/index.js"],
                ["Required env vars", "PORT, DATABASE_URL, JWT_SECRET, CLIENT_ORIGIN"],
            ],
            bullets=[
                "The backend can serve the React build in production.",
                "The repo supports workspace-based local development through root-level npm scripts.",
                "The database schema is synchronized with Prisma db push and seeded with demo accounts and shifts.",
            ],
        ),
        Section(
            "15. Demo Data and Credentials",
            table=[
                ["Account Type", "Email", "Password"],
                ["Admin", "admin@driverops.dev", "Admin@123"],
                ["Driver", "driver1@driverops.dev", "Driver@123"],
                ["Driver", "driver2@driverops.dev", "Driver@123"],
            ],
            paragraphs=[
                "The seed script also creates multiple upcoming shifts with realistic titles, locations, capacities, and notes so the application can be demonstrated immediately without manual setup.",
            ],
        ),
        Section(
            "16. Security Considerations",
            bullets=[
                "Passwords are hashed before storage.",
                "All protected endpoints require a valid bearer token.",
                "Role-based route restrictions prevent cross-role access.",
                "Zod validation reduces malformed or unsafe request payloads.",
                "Helmet is used for baseline HTTP hardening.",
                "The current demo stores JWTs in localStorage, which is acceptable for demo scope but could be upgraded to HTTP-only cookies for stronger production security.",
            ],
        ),
        Section(
            "17. Known Limitations and Recommended Enhancements",
            table=[
                ["Current Limitation", "Recommended Enhancement"],
                ["No password reset flow", "Add account recovery and reset endpoints"],
                ["No admin UI for shift editing", "Expose full shift management in frontend"],
                ["No pagination on dashboards", "Add filtering, pagination, and search"],
                ["Images stored in database", "Move to object storage for high-volume production"],
                ["No audit log", "Add activity history and review tracing"],
                ["No notifications", "Add email, SMS, or WhatsApp notifications"],
            ],
        ),
        Section(
            "18. Interview Presentation Summary",
            paragraphs=[
                "This project is a strong interview portfolio piece because it demonstrates full-stack ownership, business workflow modeling, relational data design, authentication, validation, file upload handling, and production-style deployment structure.",
                "A concise verbal summary is: I built an internal driver operations platform where drivers can securely claim and execute shifts, submit vehicle inspections with photo evidence, and admins can manage coverage and review compliance through a dedicated operational dashboard.",
            ],
        ),
    ]


def main() -> None:
    DOCS_DIR.mkdir(exist_ok=True)
    sections = get_sections()
    docx_output = build_docx(sections)
    pdf_output = build_pdf(sections)
    print(f"Created: {docx_output}")
    print(f"Created: {pdf_output}")


if __name__ == "__main__":
    main()
